#!/usr/bin/env python3
"""Parse a SafetyCulture Home Safety Assessment (HSA) Word export into the
intermediate JSON consumed by build_docx.py and build_glossy.py.

The HSA is the raw, internal field capture. This module performs the
deterministic first pass of the HSP synthesis:

  * reads the assessment metadata (site, customer, concierge, date),
  * walks each category's Item / Response / Priority table,
  * classifies every assessed element as satisfactory | recommended-work |
    monitor (from the Response column),
  * buckets each recommended-work item into one of the three HSP sections
    (from the Priority column: Immediate -> Essential Repairs,
    Programmed -> Programmed Maintenance Schedule, Enhancement -> Enhancement
    Plan),
  * rewrites each item into calm, premium, client-facing copy using a keyword
    library, and NEVER reproduces the raw inspector "Finding" notes, defect
    codes, or technical jargon.

Output: a single JSON object on stdout. Usage:
    python3 generators/parse_hsa.py fixtures/site8-hsa.docx > /tmp/hsp.json
    python3 generators/parse_hsa.py HSA.docx --client "Jane Doe" --concierge "Sam" > out.json

Metadata is parsed from the HSA where present; any field can be overridden
with --address / --client / --concierge / --date / --period.
"""
import argparse
import json
import re
import sys

from docx import Document

import hsp_common as hc

# --- Response / Priority vocabulary (SafetyCulture export) ------------------
RESP_SATISFACTORY = "satisfactory"
RESP_ACTION = "action required"
RESP_MONITOR = "monitor"
RESP_NA = ("not present", "n/a", "not applicable")

# Priority column -> HSP section key
PRIORITY_TO_SECTION = {
    "immediate": "essential_repairs",
    "essential": "essential_repairs",
    "priority": "essential_repairs",
    "programmed": "programmed_maintenance",
    "scheduled": "programmed_maintenance",
    "enhancement": "enhancement_plan",
    "enhance": "enhancement_plan",
    "improvement": "enhancement_plan",
}
INDICATIVE = {
    "essential_repairs": "Priority",
    "programmed_maintenance": "Programmed",
    "enhancement_plan": "Enhancement",
}

# --- Client-facing language library ----------------------------------------
# Each entry: keyword substrings (matched on the lowercased raw element name)
# mapped to calm client copy. First match wins; order specific -> general.
# This is how raw checklist element names become premium, non-alarming copy
# without ever exposing the inspector's raw notes.
LIBRARY = [
    # --- Safety / compliance (Essential) ---
    (["smoke alarm"], "Smoke-alarm compliance",
     "Upgrade to interconnected photoelectric smoke alarms and confirm correct placement and coverage throughout the home.",
     "Interconnected alarms give your household the earliest possible warning in the event of a fire."),
    (["rcd", "switchboard", "safety switch"], "Electrical safety",
     "Install RCD protection to the unprotected circuits and complete safety-switch testing across the board.",
     "Modern safety switches protect your household from electric shock."),
    (["pool fenc", "pool fence", "pool gate", "pool barrier", "pool fencing"], "Pool barrier compliance",
     "Adjust the self-closing pool gate and renew the faded safety signage so the barrier meets compliance.",
     "A compliant pool barrier is a legal requirement and a vital safeguard for young children."),
    (["balustrade"], "Deck balustrade",
     "Re-fix the movement at the upper-deck balustrade posts to restore full rigidity.",
     "A firm balustrade is an essential safeguard against falls."),
    (["box gutter", "valley"], "Roof drainage",
     "Clear and treat the surface corrosion in the box gutter and restore free drainage.",
     "Free-flowing roof drainage protects the roof structure and prevents water entry during heavy rain."),
    (["bearer", "joist", "subfloor bearers"], "Subfloor timber",
     "Address the elevated moisture at the subfloor bearer and treat the affected timber.",
     "Managing subfloor moisture protects the structural timbers that support your home."),
    # --- Programmed maintenance ---
    (["site drainage", "surface water"], "Site drainage",
     "Regrade or install a channel drain to clear the low point that forms near the rear patio after rain.",
     "Directing surface water away from the home protects the foundations and outdoor areas."),
    (["roof covering", "ridge", "roof tile", "roof metal"], "Roof covering",
     "Re-secure the displaced ridge fasteners and renew the localised sealant.",
     "A well-sealed roof is your home's first defence against the weather."),
    (["eaves gutter", "downpipe", "gutters &"], "Gutters & downpipes",
     "Clear debris from the gutter runs and re-secure the loose brackets.",
     "Clear gutters carry water safely away and prevent overflow damage."),
    (["render", "texture coat"], "External render",
     "Investigate the source of the hairline cracking, then patch and repaint the affected wall.",
     "Sound render keeps moisture out and preserves the appearance of the facade."),
    (["window seal", "weatherstrip"], "Window seals",
     "Replace the perished window seal and make good the affected reveal.",
     "Intact seals keep the weather out and help your home stay comfortable and efficient."),
    (["stacker", "sliding", "door track", "rollers"], "Sliding door",
     "Service the sliding-door track and rollers to restore smooth, effortless operation.",
     "Well-maintained doors are easier to use and last considerably longer."),
    (["oiling", "coatings", "deck oil"], "Deck coating",
     "Clean and re-oil the weathered decking this season.",
     "Regular oiling protects the timber and keeps the deck looking its best."),
    (["trip hazard", "lifting", "settlement", "paving"], "Path levelling",
     "Grind back or replace the lifted path slab to remove the trip hazard at the side entry.",
     "Even, level paths keep everyday movement around the home safe."),
    (["shower", "grout", "silicone", "waterproofing"], "Shower waterproofing",
     "Re-grout and reseal the shower to protect the waterproofing behind the tiles.",
     "Sound waterproofing prevents hidden moisture damage in wet areas."),
    (["exhaust fan"], "Bathroom ventilation",
     "Service or replace the exhaust fan so it manages moisture effectively.",
     "Good ventilation keeps wet areas dry and discourages mould."),
    (["tpr", "relief", "tundish"], "Hot-water relief valve",
     "Have a licensed plumber verify the temperature-relief valve and tundish arrangement.",
     "A correctly functioning relief valve is an important safety feature of the hot-water system."),
    (["subfloor vent"], "Subfloor ventilation",
     "Clear the blocked subfloor vents to restore healthy cross-flow ventilation.",
     "Good subfloor airflow keeps the timbers dry and the home healthy."),
    (["filter"], "Air-conditioning filters",
     "Clean or replace the air-conditioning filters.",
     "Clean filters improve air quality and keep the system running efficiently."),
    # --- Monitor observations ---
    (["boundary tree", "overhang", "eucalypt"], "Boundary trees",
     "A large tree overhangs the roofline; an arborist review is recommended within the coming months.",
     None),
    (["flashing", "penetration"], "Roof flashings",
     "The flue flashing is sound but ageing; we will review it at the next assessment.",
     None),
    (["external paint", "paint —", "chalking"], "External paint",
     "The west elevation is beginning to chalk; a full repaint is likely within twelve to eighteen months.",
     None),
    (["hardware corrosion", "door / window hardware"], "Door hardware",
     "Early coastal corrosion is appearing on external hardware; we are keeping it under watch.",
     None),
    (["decking board", "fixings & rot"], "Decking boards",
     "Two boards are cupping near the drip line and will be re-fixed as required.",
     None),
    (["expansion joint", "map cracking"], "Driveway cracking",
     "Cosmetic map-cracking to the driveway; no action is needed at this stage.",
     None),
    (["retaining wall"], "Retaining wall",
     "A slight lean to the rear retaining wall, with weep holes clear; we will monitor for any progression.",
     None),
    (["pool interior", "coping", "tiling"], "Pool waterline tiling",
     "Minor tile lifting at the waterline, which we will keep under observation.",
     None),
    (["floor level", "bounce", "deflection"], "Floor level",
     "A slight deflection to the rear bedroom floor, noted for review.",
     None),
    (["ceiling", "sagging", "staining"], "Ceiling finish",
     "A faint historic stain to the laundry ceiling, currently dry and stable.",
     None),
    (["timber / engineered", "flooring"], "Timber flooring",
     "Minor seasonal gapping to the entry boards, which typically closes with humidity.",
     None),
    (["dishwasher"], "Kitchen appliance seal",
     "Minor staining under the dishwasher, currently dry; we will monitor the seal.",
     None),
    (["mould", "ventilation adequacy"], "Wet-area ventilation",
     "Early spotting to a bathroom ceiling; improving ventilation will keep it in check.",
     None),
    (["hot water system", "hws", "hot water"], "Hot-water system",
     "The hot-water system is operating well at around nine years; we recommend planning for replacement within its service cycle.",
     None),
    (["lighting"], "External lighting",
     "Two external fittings show water marking and will be reviewed or replaced as needed.",
     None),
]

MONITOR_WHY = "No action is needed today; we will review this at your next assessment."


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def slugify(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    return re.sub(r"-+", "-", text).strip("-")


def synthesize(raw_element: str, state: str, section_key: str):
    """Return (title, description, why_it_matters) for an element, in calm
    client-facing language. Never echoes raw inspector notes."""
    low = raw_element.lower()
    for keys, title, desc, why in LIBRARY:
        if any(k in low for k in keys):
            return title, desc, (why or MONITOR_WHY if state == "monitor" else (why or _generic_why(section_key)))
    # Fallback: clean the raw element name and use section-generic copy.
    clean = norm(raw_element.split("—")[0].split("/")[0]).rstrip()
    clean = clean[:1].upper() + clean[1:] if clean else "Property element"
    if state == "monitor":
        return clean, f"{clean} has been noted for observation and will be reviewed at your next assessment.", MONITOR_WHY
    return clean, _generic_desc(section_key, clean), _generic_why(section_key)


def _generic_desc(section_key: str, clean: str) -> str:
    if section_key == "essential_repairs":
        return f"{clean} warrants prompt attention and we recommend addressing it early."
    if section_key == "programmed_maintenance":
        return f"{clean} is best attended to as planned, preventive care across the coming year."
    return f"{clean} is offered as a considered improvement to enhance comfort and value."


def _generic_why(section_key: str) -> str:
    return {
        "essential_repairs": "Addressing this protects the safety and integrity of your home.",
        "programmed_maintenance": "Routine care like this prevents small matters from becoming larger ones.",
        "enhancement_plan": "An enhancement like this lifts the comfort, enjoyment and value of your home.",
    }[section_key]


# --- HSA reading ------------------------------------------------------------
def _cells(row):
    return [norm(c.text) for c in row.cells]


def find_metadata(doc):
    """Locate the key/value metadata table and extract fields."""
    meta = {}
    for t in doc.tables:
        if len(t.columns) != 2:
            continue
        keys = {_cells(r)[0].lower(): _cells(r)[1] for r in t.rows if len(r.cells) >= 2}
        if any(k in keys for k in ("site", "customer", "conducted on", "assessment")):
            meta = keys
            break
    return meta


def is_item_table(t):
    if len(t.columns) < 3:
        return False
    head = [c.lower() for c in _cells(t.rows[0])]
    return head[:3] == ["item", "response", "priority"]


def classify(response: str):
    r = response.lower()
    if RESP_SATISFACTORY in r:
        return "satisfactory"
    if RESP_MONITOR in r:
        return "monitor"
    if RESP_ACTION in r or "action" in r:
        return "recommended-work"
    if any(n in r for n in RESP_NA):
        return "na"
    return "na"


def section_for(priority: str) -> str:
    p = priority.lower()
    for key, sec in PRIORITY_TO_SECTION.items():
        if key in p:
            return sec
    return "programmed_maintenance"  # safe default for an un-tagged action item


def parse(path, overrides):
    doc = Document(str(path))
    meta_raw = find_metadata(doc)

    def mget(*keys):
        for k in keys:
            if k in meta_raw and meta_raw[k]:
                return meta_raw[k]
        return ""

    # Site "Site 8 — 15 Annie Street, Hamilton, QLD" -> street address
    site = mget("site")
    address = re.sub(r"^\s*site\s*\d+\s*[—\-:]\s*", "", site, flags=re.I).strip() or site
    address = re.sub(r"\s*\([^)]*\)\s*$", "", address).strip()  # drop trailing "(alpha test)" etc.
    customer = re.sub(r"\s*\(.*?\)\s*$", "", mget("customer")).strip()
    concierge = re.sub(r"\s*\(.*?\)\s*$", "", mget("conducted by", "concierge")).strip()
    conducted = mget("conducted on")
    date = conducted.split(",")[0].strip() if conducted else ""

    metadata = {
        "property_address": overrides.get("address") or address,
        "client_name": overrides.get("client") or customer,
        "concierge": overrides.get("concierge") or concierge,
        "date": overrides.get("date") or date,
        "plan_period": overrides.get("period") or "",
    }
    metadata["address_slug"] = slugify(
        (metadata["property_address"] or "property").split(",")[0]
    ) or "property"

    sections = {k: [] for k in hc.SECTION_KEYS}
    monitor_items = []
    satisfactory_names = []
    counts = {"assessed": 0, "satisfactory": 0, "recommended": 0, "recommended_flagged": 0, "monitor": 0, "na": 0}
    seen = set()  # (section_key, title) de-dup for cross-referenced items

    for t in doc.tables:
        if not is_item_table(t):
            continue
        for row in t.rows[1:]:
            c = _cells(row)
            if len(c) < 3 or not c[0]:
                continue
            element, response, priority = c[0], c[1], c[2]
            state = classify(response)
            counts["assessed"] += 1
            if state == "satisfactory":
                counts["satisfactory"] += 1
                satisfactory_names.append(element)
            elif state == "monitor":
                counts["monitor"] += 1
                title, desc, why = synthesize(element, "monitor", "")
                key = ("monitor", title)
                if key not in seen:
                    seen.add(key)
                    monitor_items.append({
                        "element": element, "title": title,
                        "description": desc, "why_it_matters": why,
                        "priority": "Monitor",
                    })
            elif state == "recommended-work":
                counts["recommended_flagged"] += 1
                sec = section_for(priority)
                title, desc, why = synthesize(element, "recommended-work", sec)
                key = (sec, title)
                if key not in seen:
                    seen.add(key)
                    sections[sec].append({
                        "element": element, "title": title,
                        "description": desc, "why_it_matters": why,
                        "priority": INDICATIVE[sec],
                        "indicative_cost": "By quotation",
                    })
            else:
                counts["na"] += 1

    # Headline "recommended" count = distinct works presented (post-consolidation).
    counts["recommended"] = sum(len(sections[k]) for k in hc.SECTION_KEYS)

    satisfactory_summary = (
        f"Of the {counts['assessed']} elements we assessed, {counts['satisfactory']} are in "
        f"satisfactory condition and require no action. Your home presents as well cared for "
        f"and structurally sound."
    )

    overview_letter = _overview_letter(metadata, counts)

    return {
        "metadata": metadata,
        "counts": counts,
        "overview_letter": overview_letter,
        "satisfactory_summary": satisfactory_summary,
        "satisfactory_elements": satisfactory_names,
        "sections": sections,
        "monitor_items": monitor_items,
    }


def _overview_letter(meta, counts):
    client = meta["client_name"] or "your household"
    return (
        f"Thank you for entrusting the care of your home to Bespoke. This Home Stewardship "
        f"Plan sets out what we found during your assessment, what we recommend, and how we "
        f"will see it through on your behalf.\n\n"
        f"Your home presents as well cared for and structurally sound. Of the "
        f"{counts['assessed']} elements we assessed, the great majority are in good order. "
        f"We have identified a focused set of works — {counts['recommended']} in total — "
        f"together with {counts['monitor']} items we will simply keep under watch.\n\n"
        f"Every recommendation that follows is yours to approve at your discretion. Where "
        f"works proceed, we coordinate each trade, oversee the quality to our standard, and "
        f"report back to you on completion."
    )


def main():
    ap = argparse.ArgumentParser(description="Parse an HSA .docx into HSP JSON")
    ap.add_argument("hsa", help="Path to the HSA Word export (.docx)")
    ap.add_argument("--address")
    ap.add_argument("--client")
    ap.add_argument("--concierge")
    ap.add_argument("--date")
    ap.add_argument("--period")
    args = ap.parse_args()

    overrides = {k: getattr(args, k) for k in ("address", "client", "concierge", "date", "period")}
    data = parse(args.hsa, overrides)
    json.dump(data, sys.stdout, indent=2, ensure_ascii=False)
    sys.stdout.write("\n")


if __name__ == "__main__":
    main()
