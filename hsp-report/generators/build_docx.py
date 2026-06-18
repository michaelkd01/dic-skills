#!/usr/bin/env python3
"""Build the branded Home Stewardship Plan Word document from parsed HSP JSON.

    python3 generators/build_docx.py /tmp/hsp.json [--out DIR]

Emits HSP-<address-slug>.docx in the current directory (or --out DIR). Follows
the HSP document structure and brand tokens, and embeds the bundled brand
fonts so the file is self-contained.

Structure: Cover -> Overview letter -> Property at a glance -> Essential
Repairs -> Programmed Maintenance Schedule -> Enhancement Plan -> Items to
Monitor -> Closing. Empty sections still render their heading and an
"No items at this time." note.
"""
import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import hsp_common as hc
from font_embed import embed_fonts

CHARCOAL = RGBColor(0x18, 0x18, 0x18)
PASTEL = RGBColor(0xF9, 0xF8, 0xF7)
LATTE = RGBColor(0x98, 0x7F, 0x6A)
OAT = RGBColor(0xD9, 0xD1, 0xC7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

DISPLAY = "Aviano Serif Light"
BODY = "Familjen Grotesk"
BODY_MED = "Familjen Grotesk Medium"
BODY_SEMI = "Familjen Grotesk SemiBold"
TRACK_EM = -0.130  # brand display kerning token (-130 / 1000 em)


# --- low-level helpers ------------------------------------------------------
def _set(el, tag, **attrs):
    child = OxmlElement(tag)
    for k, v in attrs.items():
        child.set(qn(k), str(v))
    el.append(child)
    return child


def shade(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    _set(tcPr, "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": hex6})


def exact_row_height(row, twips):
    trPr = row._tr.get_or_add_trPr()
    _set(trPr, "w:trHeight", **{"w:val": twips, "w:hRule": "exact"})


def no_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    m = _set(tcPr, "w:tcMar")
    for side in ("top", "bottom", "start", "end"):
        _set(m, f"w:{side}", **{"w:w": 0, "w:type": "dxa"})


def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = _set(tblPr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set(borders, f"w:{edge}", **{"w:val": "none", "w:sz": 0, "w:space": 0})


def bottom_border(paragraph, hex6=hc.OAT.lstrip("#"), sz=6):
    pPr = paragraph._p.get_or_add_pPr()
    borders = _set(pPr, "w:pBdr")
    _set(borders, "w:bottom", **{"w:val": "single", "w:sz": sz, "w:space": 6, "w:color": hex6})


def style_run(run, font=BODY, size=10.5, color=CHARCOAL, bold=False, track=False,
              caps=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    # ensure east-asian / hAnsi map to the same face
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = _set(rPr, "w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(a), font)
    if track:
        # brand display tracking: -0.130 em -> twips at this size
        _set(rPr, "w:spacing", **{"w:val": int(round(TRACK_EM * size * 20))})
    if caps:
        _set(rPr, "w:caps", **{"w:val": "true"})
    return run


def para(container, text="", **kw):
    p = container.add_paragraph()
    if text:
        style_run(p.add_run(text), **kw)
    return p


def set_page_background(doc, hex6):
    """Best-effort pastel page background (Word honours it with the
    'display background' setting; LibreOffice renders it directly)."""
    body = doc.element.body
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex6)
    doc.element.insert(0, bg)
    settings = doc.settings.element
    _set(settings, "w:displayBackgroundShape")


def configure_section(section, margin_cm, a4=True):
    if a4:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    section.header_distance = Cm(min(margin_cm, 1.0))
    section.footer_distance = Cm(min(margin_cm, 1.0))


# --- page builders ----------------------------------------------------------
def add_full_charcoal_page(doc, logo_png, lines):
    """A full-bleed charcoal panel (cover / closing) with the white logo and
    centred white text lines. `lines` = list of (text, size, font, track, caps)."""
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    table.autofit = False
    table.columns[0].width = Cm(21.0)
    cell = table.cell(0, 0)
    cell.width = Cm(21.0)
    shade(cell, hc.CHARCOAL.lstrip("#"))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    exact_row_height(table.rows[0], 15800)

    # clear default empty paragraph
    cell.paragraphs[0].text = ""
    logo_p = cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(str(logo_png), width=Cm(9.0))
    logo_p.paragraph_format.space_after = Pt(28)

    for text, size, font, track, caps in lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        style_run(p.add_run(text), font=font, size=size, color=WHITE, track=track, caps=caps)
    return table


def add_display_heading(doc, text, size=22, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    style_run(p.add_run(text), font=DISPLAY, size=size, color=CHARCOAL, track=True)
    return p


def add_section_label(doc, kicker):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(kicker), font=BODY_SEMI, size=8.5, color=LATTE, caps=True, track=True)
    return p


def add_glance(doc, counts):
    add_section_label(doc, "Property at a glance")
    add_display_heading(doc, "Your home, in summary", size=20, space_before=2)
    stats = [
        (counts["assessed"], "elements assessed"),
        (counts["satisfactory"], "in satisfactory condition"),
        (counts["recommended"], "works recommended"),
        (counts["monitor"], "items to monitor"),
    ]
    table = doc.add_table(rows=1, cols=4)
    remove_table_borders(table)
    table.autofit = False
    for i, (num, label) in enumerate(stats):
        cell = table.cell(0, i)
        cell.width = Cm(4.2)
        shade(cell, hc.PASTEL.lstrip("#"))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        np = cell.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np.paragraph_format.space_before = Pt(10)
        np.paragraph_format.space_after = Pt(2)
        style_run(np.add_run(str(num)), font=DISPLAY, size=30, color=CHARCOAL, track=True)
        lp = cell.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_after = Pt(10)
        style_run(lp.add_run(label), font=BODY, size=8.5, color=LATTE)


def add_item(doc, item, monitor=False):
    # Title + indicative priority chip
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(10)
    tp.paragraph_format.space_after = Pt(2)
    style_run(tp.add_run(item["title"]), font=BODY_SEMI, size=12, color=CHARCOAL)
    chip = "  ·  " + (item.get("priority", "") or "")
    if not monitor:
        chip += "  ·  " + item.get("indicative_cost", "By quotation")
    style_run(tp.add_run(chip), font=BODY, size=9, color=LATTE)
    bottom_border(tp)

    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(2)
    style_run(dp.add_run(item["description"]), font=BODY, size=10.5, color=CHARCOAL)

    if item.get("why_it_matters"):
        wp = doc.add_paragraph()
        wp.paragraph_format.space_after = Pt(4)
        style_run(wp.add_run("Why it matters  "), font=BODY_SEMI, size=9, color=LATTE, caps=True)
        style_run(wp.add_run(item["why_it_matters"]), font=BODY, size=9.5, color=LATTE)


def add_recommendation_section(doc, sec_def, items):
    add_section_label(doc, "Recommended works")
    add_display_heading(doc, sec_def["title"], size=20, space_before=4)
    para(doc, sec_def["blurb"], font=BODY, size=10.5, color=CHARCOAL)
    if not items:
        np = doc.add_paragraph()
        np.paragraph_format.space_before = Pt(6)
        style_run(np.add_run(hc.EMPTY_SECTION_NOTE), font=BODY, size=10.5, color=LATTE)
        return
    for it in items:
        add_item(doc, it)


def build(data, out_dir):
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = CHARCOAL
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = _set(rpr, "w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), BODY)

    set_page_background(doc, hc.PASTEL.lstrip("#"))

    meta = data["metadata"]
    counts = data["counts"]

    # --- Cover (section 0) ---
    period = meta.get("plan_period") or ""
    cover_lines = [
        ("HOME STEWARDSHIP PLAN", 11, BODY_SEMI, True, True),
        (meta.get("property_address", ""), 18, DISPLAY, True, False),
    ]
    if meta.get("client_name"):
        cover_lines.append(("Prepared for " + meta["client_name"], 11, BODY, False, False))
    if meta.get("date"):
        cover_lines.append(("Assessed " + meta["date"], 10, BODY, False, False))
    if period:
        cover_lines.append((period, 10, BODY, False, False))
    add_full_charcoal_page(doc, hc.LOGO_WHITE_PNG, cover_lines)

    doc.add_section(WD_SECTION.NEW_PAGE)  # -> content section

    # small charcoal logo header on content
    hp = doc.add_paragraph()
    hp.add_run().add_picture(str(hc.LOGO_CHARCOAL_PNG), width=Cm(5.0))
    hp.paragraph_format.space_after = Pt(10)

    # --- Overview letter ---
    add_section_label(doc, "Introduction")
    add_display_heading(doc, "Your Home Stewardship Plan", size=22, space_before=2)
    for block in data["overview_letter"].split("\n\n"):
        p = para(doc, block.strip(), font=BODY, size=10.5, color=CHARCOAL)
        p.paragraph_format.space_after = Pt(8)
    if meta.get("concierge"):
        sig = para(doc, meta["concierge"], font=BODY_SEMI, size=10.5, color=CHARCOAL)
        sig.paragraph_format.space_before = Pt(6)
        para(doc, "Your Concierge, Bespoke Property Concierge", font=BODY, size=9.5, color=LATTE)

    # --- Property at a glance ---
    add_glance(doc, counts)
    gp = para(doc, data["satisfactory_summary"], font=BODY, size=10.5, color=CHARCOAL)
    gp.paragraph_format.space_before = Pt(8)

    # --- Three recommendation sections ---
    for sec_def in hc.SECTIONS:
        add_recommendation_section(doc, sec_def, data["sections"].get(sec_def["key"], []))

    # --- Items to monitor ---
    add_section_label(doc, "Under watch")
    add_display_heading(doc, "Items to Monitor", size=20)
    para(doc, "These items need no action today. We note them so nothing escapes "
              "attention, and we will review each at your next assessment.",
         font=BODY, size=10.5, color=CHARCOAL)
    if data["monitor_items"]:
        for it in data["monitor_items"]:
            add_item(doc, it, monitor=True)
    else:
        np = doc.add_paragraph()
        style_run(np.add_run(hc.EMPTY_SECTION_NOTE), font=BODY, size=10.5, color=LATTE)

    # --- Closing (section 2) ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    closing_lines = [
        ("YOUR CONCIERGE", 11, BODY_SEMI, True, True),
        ("Your home is in careful hands.", 17, DISPLAY, True, False),
        (hc.CONTACT_LINE, 10, BODY, False, False),
        ("Meticulous care for homes that deserve perfection.", 10, BODY, False, False),
    ]
    add_full_charcoal_page(doc, hc.LOGO_WHITE_PNG, closing_lines)

    # margins: sections[0]=cover, [1]=content, [2]=closing
    secs = doc.sections
    configure_section(secs[0], 0.0)
    configure_section(secs[1], 2.0)
    configure_section(secs[2], 0.0)

    slug = meta.get("address_slug") or "property"
    out_path = Path(out_dir) / f"HSP-{slug}.docx"
    doc.save(str(out_path))

    # Embed brand fonts (best-effort; document already renders via font names).
    embedded = []
    try:
        embedded = embed_fonts(out_path, [
            {"name": "Familjen Grotesk",
             "regular": hc.FONT_FILES["familjen_400"], "bold": hc.FONT_FILES["familjen_700"]},
            {"name": "Familjen Grotesk Medium", "regular": hc.FONT_FILES["familjen_500"]},
            {"name": "Familjen Grotesk SemiBold", "regular": hc.FONT_FILES["familjen_600"]},
            {"name": "Aviano Serif Light", "regular": hc.FONT_FILES["aviano_light"]},
            {"name": "Aviano Serif", "regular": hc.FONT_FILES["aviano_regular"]},
        ])
        # validate the package still opens cleanly
        Document(str(out_path))
    except Exception as exc:  # pragma: no cover - safety net
        print(f"[warn] font embedding skipped ({exc}); fonts referenced by name",
              file=sys.stderr)

    return out_path, embedded


def main():
    ap = argparse.ArgumentParser(description="Build the HSP Word document")
    ap.add_argument("json", help="parsed HSP JSON (from parse_hsa.py)")
    ap.add_argument("--out", default=".", help="output directory")
    args = ap.parse_args()
    data = json.load(open(args.json, encoding="utf-8"))
    out_path, embedded = build(data, args.out)
    print(f"Wrote {out_path}")
    if embedded:
        print(f"Embedded fonts: {len(embedded)} variants")


if __name__ == "__main__":
    main()
